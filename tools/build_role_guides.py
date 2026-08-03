from __future__ import annotations

import sys
from pathlib import Path

RUNTIME = Path.home() / "AppData" / "Local" / "Temp" / "codex-docx-runtime"
SKILL_ROOT = Path.home() / ".codex" / "plugins" / "cache" / "openai-primary-runtime" / "documents" / "26.521.10419" / "skills" / "documents"
sys.path.insert(0, str(RUNTIME))
sys.path.insert(0, str(SKILL_ROOT / "scripts"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from table_geometry import apply_table_geometry, audit_docx_tables


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "docs" / "panduan"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}

COLORS = {
    "brand_dark": "3B2418",
    "brand": "6B3F2A",
    "accent": "B7653F",
    "accent_soft": "F0D9C9",
    "cream": "F6F1E8",
    "surface": "FFFDF9",
    "border": "D9CCBC",
    "text": "2A211C",
    "muted": "6F6259",
    "success": "2F6B4F",
    "warning": "9A6717",
    "danger": "A33C32",
    "info": "3F6275",
    "table_fill": "E8EEF5",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9CCBC", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_numbering_definition(document, kind, marker):
    numbering = document.part.numbering_part.element
    existing_abs = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    existing_nums = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_nums, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), kind)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text_node, end])
    set_run_font(run, size=9, color=COLORS["muted"])


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["text"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, COLORS["brand"], 18, 10),
        "Heading 2": (13, COLORS["brand"], 14, 7),
        "Heading 3": (12, COLORS["brand_dark"], 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles.add_style("Manual Title", WD_STYLE_TYPE.PARAGRAPH)
    title.font.name = "Georgia"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(COLORS["brand_dark"])
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True

    subtitle = styles.add_style("Manual Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(COLORS["brand"])
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.keep_with_next = True

    kicker = styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    kicker.font.name = "Calibri"
    kicker._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    kicker._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    kicker.font.size = Pt(9)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(COLORS["accent"])
    kicker.paragraph_format.space_after = Pt(8)

    small = styles.add_style("Small Muted", WD_STYLE_TYPE.PARAGRAPH)
    small.font.name = "Calibri"
    small._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    small._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    small.font.size = Pt(9)
    small.font.color.rgb = rgb(COLORS["muted"])
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.15


def configure_section(section, role_label):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"SISTEM PENJURIAN KARNAVAL  |  {role_label.upper()}")
    set_run_font(run, size=8.5, color=COLORS["muted"], bold=True)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.35), WD_TAB_ALIGNMENT.RIGHT)
    footer_run = footer_paragraph.add_run("Panduan Penggunaan - Randuagung 2026")
    set_run_font(footer_run, size=8.5, color=COLORS["muted"])
    footer_paragraph.add_run("\t")
    add_page_field(footer_paragraph)


def new_document(role_label):
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0], role_label)
    document._bullet_num_id = add_numbering_definition(document, "bullet", "•")
    document._number_num_id = add_numbering_definition(document, "decimal", "%1.")
    document._check_num_id = add_numbering_definition(document, "bullet", "☐")
    props = document.core_properties
    props.title = f"Panduan Penggunaan - {role_label}"
    props.subject = "Sistem Penjurian Karnaval Kecamatan Randuagung 2026"
    props.author = "Panitia Karnaval Kecamatan Randuagung"
    props.keywords = "panduan, karnaval, penjurian, Randuagung"
    props.comments = "Panduan operasional berbasis implementasi aplikasi versi Agustus 2026."
    return document


def add_cover(document, role_label, tagline):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(86)
    kicker = document.add_paragraph(style="Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.add_run("BUKU PANDUAN PENGGUNA")
    title = document.add_paragraph(style="Manual Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(role_label)
    subtitle = document.add_paragraph(style="Manual Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Sistem Penjurian Karnaval Kecamatan Randuagung 2026")

    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(24)
    run = line.add_run(tagline)
    set_run_font(run, size=11.5, color=COLORS["muted"], italic=True)

    identity = document.add_paragraph()
    identity.alignment = WD_ALIGN_PARAGRAPH.CENTER
    identity.paragraph_format.space_before = Pt(52)
    identity.paragraph_format.space_after = Pt(4)
    run = identity.add_run("PANITIA KARNAVAL KECAMATAN RANDUAGUNG")
    set_run_font(run, size=10, color=COLORS["brand"], bold=True)
    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_after = Pt(0)
    run = version.add_run("Versi 1.0  |  Agustus 2026")
    set_run_font(run, size=9.5, color=COLORS["muted"])
    document.add_page_break()


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    return paragraph


def add_body(document, text, bold_lead=None):
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        first.bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_list(document, items, numbered=False, checklist=False):
    num_id = document._check_num_id if checklist else document._number_num_id if numbered else document._bullet_num_id
    paragraphs = []
    for item in items:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_element])
        p_pr.append(num_pr)
        if isinstance(item, tuple):
            label, detail = item
            run = paragraph.add_run(label)
            run.bold = True
            paragraph.add_run(detail)
        else:
            paragraph.add_run(item)
        paragraphs.append(paragraph)
    return paragraphs


def add_callout(document, title, text, tone="info"):
    color = COLORS.get(tone, COLORS["info"])
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Twips(180)
    paragraph.paragraph_format.right_indent = Twips(180)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(paragraph, COLORS["cream"])
    set_paragraph_left_border(paragraph, color)
    lead = paragraph.add_run(f"{title}. ")
    set_run_font(lead, size=10.5, color=color, bold=True)
    detail = paragraph.add_run(text)
    set_run_font(detail, size=10.5, color=COLORS["text"])
    return paragraph


def add_definition(document, label, value):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    lead = paragraph.add_run(f"{label}: ")
    lead.bold = True
    paragraph.add_run(value)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, text in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, COLORS["accent_soft"])
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5, color=COLORS["brand_dark"], bold=True)
    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.5, color=COLORS["text"])
    apply_table_geometry(table, widths, table_width_dxa=CONTENT_WIDTH_DXA, indent_dxa=TABLE_INDENT_DXA, cell_margins_dxa=CELL_MARGINS)
    after = document.add_paragraph()
    after.style = document.styles["Small Muted"]
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_common_access(document, role_label, destination):
    add_heading(document, "1. Memulai", 1)
    add_body(document, f"Panduan ini ditujukan untuk pengguna dengan role {role_label}. Setelah login berhasil, sistem mengarahkan Anda ke {destination}.")
    add_heading(document, "Masuk ke sistem", 2)
    add_list(document, [
        "Buka alamat aplikasi yang diberikan panitia.",
        "Isi email dan kata sandi akun pribadi Anda pada halaman Masuk.",
        "Pilih tombol Masuk dan tunggu sampai halaman role Anda terbuka.",
        "Pastikan nama, role, dan status koneksi pada bagian atas aplikasi sesuai.",
    ], numbered=True)
    add_callout(document, "Keamanan akun", "Jangan berbagi akun, kata sandi, tautan pemulihan, atau tangkapan layar yang berisi data internal. Setiap aksi penting tercatat atas nama akun yang sedang digunakan.", "danger")
    add_heading(document, "Lupa kata sandi", 2)
    add_list(document, [
        "Isi email pada halaman Masuk.",
        "Pilih Lupa kata sandi dan periksa kotak masuk email.",
        "Buka tautan pemulihan, isi kata sandi baru minimal delapan karakter, lalu simpan.",
        "Masuk kembali menggunakan kata sandi baru.",
    ], numbered=True)
    add_heading(document, "Keluar dengan aman", 2)
    add_body(document, "Pilih tombol Keluar dari sidebar atau menu akun. Pastikan draf atau perubahan terakhir sudah berstatus tersimpan sebelum keluar, terutama pada perangkat bersama.")


def add_shared_status_help(document):
    add_heading(document, "Status koneksi dan pesan sistem", 2)
    add_table(document, ["Pesan", "Arti dan tindakan"], [
        ("Terhubung", "Aplikasi dapat berkomunikasi dengan Supabase."),
        ("Luring", "Periksa jaringan. Jangan menutup tab bila masih ada input yang belum tersimpan."),
        ("Sesi berakhir", "Masuk kembali. Data yang sudah berhasil dikirim tetap aman."),
        ("Akses tidak diizinkan", "Role Anda tidak berhak membuka halaman tersebut. Kembali ke halaman utama."),
        ("Data belum dapat dimuat", "Muat ulang satu kali. Bila tetap terjadi, catat waktu dan halaman lalu hubungi Admin."),
    ], [2100, 7260])


def build_super_admin():
    role = "Super Admin"
    doc = new_document(role)
    add_cover(doc, role, "Menjaga konfigurasi, kewenangan sensitif, dan hasil resmi tetap dapat diaudit.")
    add_common_access(doc, role, "Ringkasan Operasional")

    add_heading(doc, "2. Ruang lingkup kewenangan", 1)
    add_body(doc, "Super Admin memiliki seluruh kemampuan Admin Panitia dan menjadi satu-satunya role pada antarmuka yang dapat menangani pengecualian penilaian sensitif.")
    add_table(doc, ["Kewenangan", "Penggunaan"], [
        ("Akun dan role", "Mengelola profil termasuk akun Admin, status aktif, dan penugasan petugas."),
        ("Konfigurasi event", "Mengatur event, rubrik, aturan hasil, lalu membuka atau menutup penjurian."),
        ("Unlock nilai", "Membuka lembar submitted dengan alasan resmi agar Juri dapat mengoreksi."),
        ("Waiver", "Menyetujui nilai hilang bila minimal dua nilai submitted tersedia dan ada berita acara."),
        ("Keputusan Dewan Juri", "Mengisi prioritas tie-break beserta referensi berita acara."),
        ("Penalti berat", "Mengonfirmasi penalti yang jenisnya memerlukan persetujuan Super Admin."),
        ("Publikasi dan audit", "Menerbitkan snapshot hasil dan menelusuri seluruh jejak perubahan."),
    ], [2600, 6760])
    add_callout(doc, "Prinsip empat mata", "Untuk unlock, waiver, penalti berat, tie-break, dan publikasi, cocokkan tindakan dengan dokumen atau keputusan panitia sebelum menekan tombol konfirmasi.", "warning")

    add_heading(doc, "3. Menyiapkan akun dan penugasan", 1)
    add_heading(doc, "Buat akun Auth terlebih dahulu", 2)
    add_body(doc, "Aplikasi tidak membuat password. Buat akun pada Supabase Authentication - Users, lalu salin User ID (UUID). Jangan pernah memasukkan service-role key ke browser.")
    add_heading(doc, "Tambahkan profil", 2)
    add_list(doc, [
        "Buka Pengguna, lalu pilih Tambahkan profil.",
        "Masukkan User ID, nama lengkap petugas, dan role yang benar.",
        "Simpan profil dan pastikan statusnya Aktif.",
        "Untuk akun Admin baru, lakukan langkah ini sebagai Super Admin.",
    ], numbered=True)
    add_heading(doc, "Tetapkan petugas wajib", 2)
    add_list(doc, [
        "Tetapkan satu Juri aktif untuk masing-masing lokasi Start, B. Edi, dan Finish.",
        "Tetapkan satu Operator verifikator aktif untuk masing-masing titik Junaidi, B. Sul, dan Toko Aminah.",
        "Gunakan Operator tambahan tanpa penugasan atraksi untuk monitor status dan insiden bila diperlukan.",
        "Jangan menonaktifkan satu-satunya Super Admin aktif atau akun yang sedang digunakan.",
    ], checklist=True)

    add_heading(doc, "4. Menyiapkan event", 1)
    add_heading(doc, "Data event dan aturan hasil", 2)
    add_list(doc, [
        "Buka Pengaturan Event dan periksa nama, tanggal, zona waktu, serta rute.",
        "Pastikan delapan kriteria aktif dan total maksimum tepat 100.",
        "Pastikan metode agregasi, pembulatan, mode atraksi, dan poin atraksi sesuai keputusan resmi.",
        "Periksa tiga lokasi, kategori peserta, tiga titik atraksi, dan jenis penalti.",
    ], numbered=True)
    add_heading(doc, "Readiness sebelum membuka penjurian", 2)
    add_table(doc, ["Pemeriksaan", "Syarat"], [
        ("Event", "Nama, tanggal, dan rute terisi."),
        ("Lokasi", "Tepat tiga lokasi penilaian aktif."),
        ("Juri", "Tiga Juri aktif memiliki penugasan berbeda."),
        ("Atraksi", "Tiga titik dan tiga Operator verifikator aktif tersedia."),
        ("Rubrik", "Delapan kriteria berjumlah tepat 100."),
        ("Peserta", "Minimal satu peserta aktif dan nomor urut unik."),
        ("Aturan", "Kategori, penalti, dan tie-break tersedia."),
    ], [2500, 6860])
    add_callout(doc, "Konfigurasi terkunci", "Setelah penjurian dibuka, event, rubrik, dan aturan hasil tidak dapat diubah melalui form. Periksa ulang sebelum mengonfirmasi Buka penjurian.", "warning")

    add_heading(doc, "5. Mengelola peserta dan operasional", 1)
    add_heading(doc, "Peserta manual atau CSV", 2)
    add_body(doc, "Pada menu Peserta Anda dapat mencari, memfilter, menambah, mengedit, melihat detail, dan menonaktifkan peserta. Nomor urut harus unik. Jumlah anggota di bawah 30 wajib memiliki alasan pengecualian.")
    add_definition(doc, "Kolom CSV", "nomor_urut, nama, kategori, tema, koordinator, kontak, jumlah_anggota, jadwal_berangkat, alasan_pengecualian")
    add_list(doc, [
        "Gunakan CSV UTF-8 berukuran maksimum 2 MB.",
        "Tinjau jumlah baris pada dialog konfirmasi sebelum impor.",
        "Bila validasi gagal, perbaiki file lalu ulangi; jangan mengubah sebagian data secara acak.",
    ])
    add_heading(doc, "Hari-H", 2)
    add_list(doc, [
        "Pantau Jadwal & Status untuk perubahan perjalanan peserta.",
        "Pantau Penjurian untuk status 0/3 sampai 3/3 per peserta.",
        "Pantau Atraksi Wajib sampai tidak ada unable_to_verify.",
        "Catat penalti sebagai draf, verifikasi alasan, lalu konfirmasi atau batalkan.",
    ], checklist=True)

    add_heading(doc, "6. Tindakan sensitif penilaian", 1)
    add_heading(doc, "Membuka kembali nilai", 2)
    add_list(doc, [
        "Buka Penjurian dan pilih Kelola pada peserta.",
        "Pilih Buka kembali nilai, pilih lokasi yang memiliki status submitted, lalu isi alasan.",
        "Konfirmasi. Status menjadi unlocked dan Juri lokasi tersebut dapat memperbaiki serta submit ulang.",
        "Pantau sampai status kembali submitted dan cek Audit Log.",
    ], numbered=True)
    add_heading(doc, "Memberikan waiver", 2)
    add_list(doc, [
        "Pastikan peserta sudah memiliki sedikitnya dua nilai submitted.",
        "Pilih Waiver nilai hilang dan lokasi yang tidak dapat menyerahkan nilai.",
        "Isi alasan faktual dan referensi berita acara resmi.",
        "Konfirmasi lalu periksa label waiver pada monitor dan hasil internal.",
    ], numbered=True)
    add_callout(doc, "Waiver bukan jalan pintas", "Gunakan hanya untuk keadaan luar biasa. Waiver tanpa dua nilai submitted, alasan, atau referensi berita acara akan ditolak database.", "danger")

    add_heading(doc, "7. Finalisasi dan publikasi", 1)
    add_list(doc, [
        "Pastikan seluruh lembar submitted atau memiliki waiver yang sah.",
        "Pastikan seluruh atraksi sudah performed atau not_performed.",
        "Konfirmasi atau batalkan seluruh penalti draf.",
        "Tutup penjurian dari Pengaturan Event.",
        "Buka Hasil, periksa kategori Pendidikan dan Umum serta status incomplete.",
        "Bila tie-break meminta keputusan Dewan Juri, isi prioritas dan referensi berita acara.",
        "Ekspor CSV untuk pemeriksaan silang.",
        "Pilih Buat snapshot & terbitkan hanya setelah seluruh pihak menyetujui hasil.",
    ], numbered=True)
    add_callout(doc, "Snapshot resmi", "Snapshot yang diterbitkan menjadi sumber halaman publik /hasil. Rincian nilai per Juri dan catatan internal tidak ikut dipublikasikan.", "info")

    add_heading(doc, "8. Checklist penutupan", 1)
    add_list(doc, [
        "Status event sudah scoring_closed sebelum publikasi.",
        "Tidak ada nilai incomplete atau tie-break tanpa keputusan.",
        "Tidak ada atraksi unable_to_verify.",
        "Tidak ada penalti draf yang belum diputuskan.",
        "CSV hasil telah diperiksa silang.",
        "Snapshot hasil berhasil diterbitkan dan /hasil dapat dibuka tanpa login.",
        "Audit Log menunjukkan pelaku dan waktu tindakan sensitif.",
    ], checklist=True)
    add_shared_status_help(doc)
    return doc


def build_admin():
    role = "Admin Panitia"
    doc = new_document(role)
    add_cover(doc, role, "Mengelola data event, operasional peserta, penalti, dan publikasi hasil.")
    add_common_access(doc, role, "Ringkasan Operasional")

    add_heading(doc, "2. Peta menu", 1)
    add_table(doc, ["Menu", "Fungsi utama"], [
        ("Ringkasan", "Status event, peserta, progres nilai, atraksi, dan penalti."),
        ("Peserta", "Tambah, edit, impor CSV, detail, pencarian, filter, dan nonaktifkan."),
        ("Jadwal & Status", "Pantau urutan, waktu aktual, dan status perjalanan."),
        ("Penjurian", "Pantau kelengkapan tiga Juri tanpa melihat nilai saat scoring_open."),
        ("Atraksi Wajib", "Pantau tiga titik dan poin kepatuhan."),
        ("Penalti", "Buat draf, konfirmasi, atau batalkan penalti."),
        ("Hasil", "Preview, ekspor CSV, snapshot, dan publikasi."),
        ("Pengguna", "Kelola profil Juri, Operator, Viewer, dan penugasan."),
        ("Pengaturan Event", "Atur event, rubrik, aturan hasil, readiness, buka/tutup penjurian."),
        ("Audit Log", "Cari, filter, dan ekspor jejak perubahan."),
    ], [2500, 6860])
    add_callout(doc, "Batas role", "Admin tidak dapat mengelola akun Admin/Super Admin, unlock nilai, memberi waiver, atau mengisi keputusan Dewan Juri. Hubungi Super Admin untuk tindakan tersebut.", "warning")

    add_heading(doc, "3. Persiapan event", 1)
    add_list(doc, [
        "Periksa data event, tanggal, rute, dan aturan hasil pada Pengaturan Event.",
        "Pastikan delapan kriteria berjumlah tepat 100.",
        "Masukkan peserta dan periksa nomor urut serta kategori.",
        "Pastikan tiga Juri dan tiga verifikator atraksi aktif telah ditugaskan.",
        "Periksa seluruh indikator Kesiapan sampai berstatus Siap.",
        "Konfirmasikan Buka penjurian setelah briefing dan pemeriksaan terakhir.",
    ], checklist=True)
    add_heading(doc, "Mengelola profil petugas", 2)
    add_body(doc, "Buat akun terlebih dahulu pada Supabase Authentication. Di aplikasi, pilih Pengguna - Tambahkan profil, lalu masukkan User ID, nama lengkap, dan role Juri, Operator, atau Viewer.")
    add_list(doc, [
        "Tugaskan Juri hanya ke satu lokasi aktif.",
        "Tugaskan Operator verifikator hanya ke satu titik atraksi.",
        "Nonaktifkan akun yang tidak lagi digunakan, tetapi jangan memakai akun bersama.",
    ])

    add_heading(doc, "4. Mengelola peserta", 1)
    add_heading(doc, "Tambah atau edit manual", 2)
    add_list(doc, [
        "Buka Peserta dan pilih Tambah peserta, atau pilih ikon Edit pada baris peserta.",
        "Isi nomor urut, kategori, nama, jumlah anggota, dan data pendukung.",
        "Bila anggota kurang dari 30, isi alasan pengecualian.",
        "Simpan lalu buka Detail peserta untuk memeriksa jadwal, progres, dan riwayat status.",
    ], numbered=True)
    add_heading(doc, "Impor CSV", 2)
    add_definition(doc, "Kolom dikenali", "nomor_urut, nama, kategori, tema, koordinator, kontak, jumlah_anggota, jadwal_berangkat, alasan_pengecualian")
    add_list(doc, [
        "Gunakan file .csv UTF-8 maksimum 2 MB.",
        "Pastikan kategori tertulis Pendidikan atau Umum dan jadwal memakai HH:mm.",
        "Pilih Impor CSV, pilih file, tinjau jumlah baris, lalu konfirmasi.",
        "Jika ada pesan baris invalid, perbaiki sumber CSV sebelum mengulang.",
    ], numbered=True)
    add_callout(doc, "Nonaktifkan, bukan hapus", "Peserta yang dinonaktifkan tidak tampil lagi pada operasional dan hasil, tetapi riwayatnya tetap tersimpan untuk audit.", "info")

    add_heading(doc, "5. Operasional hari-H", 1)
    add_heading(doc, "Jadwal dan status", 2)
    add_table(doc, ["Urutan status", "Kapan digunakan"], [
        ("Terdaftar - Antre", "Peserta hadir dan menunggu giliran."),
        ("Dipanggil - Tampil", "Peserta dipanggil dan sedang tampil di titik terkait."),
        ("Berangkat - Tiba", "Peserta bergerak di rute dan tiba di lokasi berikutnya."),
        ("Selesai", "Seluruh rangkaian peserta selesai."),
        ("Bermasalah", "Ada kejadian yang perlu tindak lanjut."),
        ("Mengundurkan diri", "Peserta tidak melanjutkan keikutsertaan."),
    ], [2600, 6760])
    add_body(doc, "Pada Ubah status, periksa status baru, waktu aktual, dan catatan sebelum menyimpan. Perubahan menghasilkan status log.")
    add_heading(doc, "Monitor penjurian dan atraksi", 2)
    add_list(doc, [
        "Gunakan matriks Penjurian untuk mengejar lembar draft atau belum dimulai.",
        "Nilai Juri lain tetap tersembunyi selama penjurian terbuka.",
        "Pada Atraksi Wajib, selesaikan semua unable_to_verify sebelum finalisasi.",
        "Gunakan Segarkan bila pembaruan Realtime belum muncul.",
    ])

    add_heading(doc, "6. Penalti", 1)
    add_table(doc, ["Tingkat", "Pengurangan", "Catatan"], [
        ("Ringan", "-2", "Dapat dikonfirmasi sesuai kewenangan Admin."),
        ("Sedang", "-5", "Dapat dikonfirmasi sesuai kewenangan Admin."),
        ("Berat", "-10", "Hanya dapat dikonfirmasi oleh Super Admin."),
    ], [2200, 1800, 5360])
    add_list(doc, [
        "Pilih Catat penalti, peserta, tingkat, alasan faktual, dan waktu kejadian.",
        "Simpan sebagai draf agar dapat ditinjau terlebih dahulu.",
        "Pilih Konfirmasi jika keputusan sah; hanya penalti confirmed yang dihitung.",
        "Untuk penalti Berat, minta Super Admin meninjau dan melakukan konfirmasi.",
        "Untuk membatalkan, isi alasan. Penalti tetap ada di audit tetapi tidak dihitung.",
    ], numbered=True)
    add_callout(doc, "Periksa peserta dan alasan", "Jangan mengonfirmasi penalti berdasarkan pesan lisan yang belum diverifikasi. Nilai pengurangan langsung memengaruhi hasil akhir.", "danger")

    add_heading(doc, "7. Finalisasi hasil", 1)
    add_list(doc, [
        "Pastikan penilaian, atraksi, dan penalti telah lengkap.",
        "Jika ada nilai hilang atau tie-break, minta Super Admin menyelesaikannya.",
        "Tutup penjurian dari Pengaturan Event.",
        "Buka Hasil dan periksa Pendidikan serta Umum secara terpisah.",
        "Cocokkan nilai Juri agregat, poin atraksi, penalti, nilai akhir, dan peringkat.",
        "Ekspor CSV untuk pemeriksaan silang.",
        "Pilih Buat snapshot & terbitkan setelah hasil disetujui panitia.",
        "Buka /hasil tanpa login untuk memastikan hasil publik tampil.",
    ], numbered=True)
    add_callout(doc, "Jangan terbitkan terlalu dini", "Tombol publikasi diblokir bila event belum ditutup atau ada hasil incomplete. Snapshot yang terbit menjadi hasil resmi publik.", "warning")

    add_heading(doc, "8. Audit dan checklist akhir", 1)
    add_list(doc, [
        "Cari pelaku atau aksi pada Audit Log bila ada ketidaksesuaian.",
        "Gunakan filter Konfigurasi, Peserta, Penilaian, atau Penalti.",
        "Ekspor audit ke CSV bila diperlukan untuk berita acara.",
    ])
    add_heading(doc, "Checklist Admin", 2)
    add_list(doc, [
        "Data peserta dan jadwal sudah benar.",
        "Tiga Juri dan tiga verifikator aktif sudah bertugas.",
        "Tidak ada lembar nilai belum lengkap.",
        "Tidak ada atraksi unable_to_verify.",
        "Tidak ada penalti draf tersisa.",
        "Status event sudah ditutup sebelum publikasi.",
        "Hasil dua kategori dan halaman publik sudah diperiksa.",
    ], checklist=True)
    add_shared_status_help(doc)
    return doc


def build_judge():
    role = "Juri"
    doc = new_document(role)
    add_cover(doc, role, "Menilai secara independen, konsisten, dan aman pada lokasi penugasan.")
    add_common_access(doc, role, "Daftar Penilaian")

    add_heading(doc, "2. Sebelum mulai menilai", 1)
    add_list(doc, [
        "Pastikan nama lokasi penugasan pada bagian atas sesuai dengan posisi Anda.",
        "Jika muncul Belum ada penugasan aktif, jangan menilai; hubungi Admin Panitia.",
        "Gunakan satu perangkat dan satu akun pribadi selama sesi bila memungkinkan.",
        "Pastikan indikator koneksi Terhubung dan baterai perangkat mencukupi.",
        "Jangan meminta atau membandingkan nilai Juri lain.",
    ], checklist=True)
    add_callout(doc, "Independensi Juri", "Nilai hanya berdasarkan penampilan yang Anda amati di titik tugas sendiri. Sistem tidak menampilkan nilai Juri lain selama penjurian berlangsung.", "info")

    add_heading(doc, "3. Rubrik penilaian", 1)
    add_table(doc, ["Kriteria", "Maksimum"], [
        ("Konsep dan Orisinalitas", "20"),
        ("Alur Cerita", "15"),
        ("Artistik Visual", "20"),
        ("Koreografi dan Penyajian", "15"),
        ("Nilai Budaya", "10"),
        ("Entertainment Value", "10"),
        ("Musik dan Tata Suara", "5"),
        ("Kedisiplinan dan Kekompakan", "5"),
        ("Total", "100"),
    ], [7460, 1900])
    add_list(doc, [
        "Gunakan angka bulat dari 0 sampai batas maksimum setiap kriteria.",
        "Semua kriteria wajib diisi sebelum nilai dikirim.",
        "Nilai 0 wajib disertai alasan pada kolom yang muncul.",
        "Catatan umum bersifat opsional dan sebaiknya faktual.",
        "Total dihitung otomatis oleh sistem; jangan menghitung atau mengetik total sendiri.",
    ])

    add_heading(doc, "4. Mengisi dan menyimpan nilai", 1)
    add_list(doc, [
        "Buka Penilaian dan pilih Lanjutkan nilai pada peserta yang sesuai nomor urut.",
        "Periksa nama peserta, kategori, tema, dan lokasi penugasan.",
        "Isi delapan kriteria. Gunakan tombol tambah/kurang atau keyboard angka.",
        "Bila memberi 0, tulis alasan yang dapat dipertanggungjawabkan.",
        "Periksa total pada ringkasan kanan atau bagian bawah layar ponsel.",
        "Tunggu status Tersimpan, atau pilih Simpan draf sebelum berpindah peserta.",
    ], numbered=True)
    add_heading(doc, "Arti status penyimpanan", 2)
    add_table(doc, ["Status", "Tindakan"], [
        ("Belum tersimpan", "Tunggu autosave atau pilih Simpan draf."),
        ("Menyimpan", "Jangan tutup tab atau keluar dari halaman."),
        ("Tersimpan", "Draf sudah tersimpan di Supabase dan perangkat."),
        ("Error", "Periksa jaringan, jangan hapus input, lalu coba Simpan draf lagi."),
    ], [2400, 6960])
    add_callout(doc, "Draf lokal", "Sistem menyimpan draf menggunakan event, lembar nilai, dan akun Juri. Jangan membersihkan data situs/browser saat masih ada draf yang belum berhasil tersimpan ke Supabase.", "warning")

    add_heading(doc, "5. Review, kirim, dan kunci", 1)
    add_list(doc, [
        "Pilih Review nilai setelah semua kriteria selesai.",
        "Jika ada pesan validasi, kembali ke kriteria yang ditandai dan perbaiki.",
        "Pada dialog review, cocokkan setiap skor dan total /100.",
        "Pilih Kirim & kunci nilai hanya jika Anda yakin.",
        "Tunggu notifikasi berhasil dan perpindahan ke Riwayat Nilai Saya.",
    ], numbered=True)
    add_callout(doc, "Setelah submit", "Lembar menjadi terkunci dan tidak dapat diubah sendiri. Koreksi hanya dapat dilakukan setelah Super Admin membuka kembali nilai dengan alasan resmi.", "danger")
    add_heading(doc, "Koreksi setelah unlock", 2)
    add_list(doc, [
        "Pastikan Admin/Super Admin memberi tahu peserta dan alasan koreksi.",
        "Buka kembali peserta yang berstatus unlocked.",
        "Perbaiki hanya bagian yang memang salah, simpan, review, lalu submit ulang.",
        "Pastikan status kembali submitted pada Riwayat Nilai Saya.",
    ], numbered=True)

    add_heading(doc, "6. Saat koneksi bermasalah", 1)
    add_list(doc, [
        "Tetap di halaman penilaian dan jangan logout.",
        "Input yang baru diubah dipertahankan sementara di perangkat.",
        "Aktifkan kembali jaringan dan tunggu indikator Terhubung.",
        "Pilih Simpan draf dan tunggu status Tersimpan.",
        "Jangan menekan Kirim berulang kali; tunggu respons sistem.",
        "Jika konflik versi muncul, hentikan pengeditan dan hubungi Admin sebelum memuat ulang.",
    ], numbered=True)
    add_callout(doc, "Submit tidak dilakukan otomatis", "Draf luring tidak pernah dianggap nilai final. Anda tetap harus melakukan Review dan Kirim & kunci setelah koneksi pulih.", "info")

    add_heading(doc, "7. Checklist Juri", 1)
    add_list(doc, [
        "Lokasi penugasan pada layar sudah benar.",
        "Nomor dan nama peserta sudah sesuai.",
        "Delapan kriteria berisi angka bulat dalam batas maksimum.",
        "Setiap nilai 0 memiliki alasan.",
        "Total sudah diperiksa.",
        "Status draf Tersimpan sebelum review.",
        "Dialog review sudah diperiksa sebelum submit.",
        "Nilai yang selesai muncul di Riwayat Nilai Saya.",
    ], checklist=True)
    add_shared_status_help(doc)
    return doc


def build_operator():
    role = "Operator Lapangan"
    doc = new_document(role)
    add_cover(doc, role, "Memperbarui status peserta, verifikasi atraksi, dan kejadian lapangan secara faktual.")
    add_common_access(doc, role, "Monitor Peserta")

    add_heading(doc, "2. Peta menu dan batas kewenangan", 1)
    add_table(doc, ["Menu", "Fungsi"], [
        ("Monitor Peserta", "Cari peserta, filter status, catat status dan waktu aktual."),
        ("Atraksi Wajib", "Verifikasi hanya titik atraksi yang ditugaskan kepada akun Anda."),
        ("Insiden", "Catat kejadian dan tandai sudah ditangani."),
    ], [2500, 6860])
    add_callout(doc, "Batas role", "Operator tidak dapat mengubah nilai Juri, membuat keputusan hasil, mengelola akun, atau menerbitkan hasil.", "warning")

    add_heading(doc, "3. Memperbarui status peserta", 1)
    add_table(doc, ["Status", "Gunakan ketika"], [
        ("Terdaftar", "Peserta tercatat tetapi belum masuk antrean."),
        ("Antre", "Peserta siap dan menunggu giliran."),
        ("Dipanggil", "Peserta dipanggil menuju area tampil."),
        ("Tampil", "Peserta sedang melakukan penampilan."),
        ("Berangkat", "Peserta meninggalkan titik menuju rute berikutnya."),
        ("Tiba", "Peserta tiba di lokasi berikutnya atau finish."),
        ("Selesai", "Rangkaian peserta telah selesai."),
        ("Bermasalah", "Ada kendala yang perlu dicatat atau ditindaklanjuti."),
        ("Mengundurkan diri", "Peserta resmi tidak melanjutkan."),
    ], [2300, 7060])
    add_list(doc, [
        "Cari peserta berdasarkan nama atau nomor urut.",
        "Pilih Ubah status pada baris yang tepat.",
        "Pilih status baru dan periksa waktu aktual yang ditampilkan.",
        "Tambahkan catatan bila ada konteks penting.",
        "Pilih Simpan status dan tunggu notifikasi berhasil.",
    ], numbered=True)
    add_callout(doc, "Pastikan peserta benar", "Sebutkan nomor urut saat berkoordinasi. Status dan waktu aktual masuk ke log dan dapat memengaruhi keputusan operasional.", "danger")

    add_heading(doc, "4. Verifikasi atraksi wajib", 1)
    add_body(doc, "Akun verifikator hanya dapat mengubah kolom pada titik penugasan sendiri. Jika muncul Belum ada penugasan verifier, hubungi Admin dan jangan memakai akun petugas lain.")
    add_table(doc, ["Status", "Arti"], [
        ("performed", "Atraksi wajib ditampilkan; pada mode fixed_points memberi poin sesuai konfigurasi."),
        ("not_performed", "Atraksi tidak ditampilkan; tidak memperoleh poin."),
        ("unable_to_verify", "Belum dapat dipastikan; tidak memberi poin sementara dan memblokir finalisasi."),
    ], [2500, 6860])
    add_list(doc, [
        "Buka Atraksi Wajib dan pastikan nama titik penugasan benar.",
        "Temukan peserta berdasarkan nomor urut.",
        "Pilih status sesuai pengamatan langsung.",
        "Tunggu notifikasi berhasil sebelum berpindah peserta.",
        "Jangan membiarkan unable_to_verify tanpa tindak lanjut sampai penutupan.",
    ], numbered=True)
    add_callout(doc, "Verifikasi faktual", "Jangan menebak status atraksi. Bila pandangan terhalang atau informasi belum cukup, gunakan unable_to_verify lalu koordinasikan verifikasi resmi.", "warning")

    add_heading(doc, "5. Mencatat insiden", 1)
    add_list(doc, [
        "Buka Insiden lalu pilih Catat insiden.",
        "Pilih peserta yang benar.",
        "Pilih jenis: Keterlambatan, Kendala rute, Ketertiban, atau Lainnya.",
        "Tulis catatan singkat, faktual, dan tanpa opini pribadi.",
        "Simpan insiden. Setelah tindak lanjut selesai, pilih Tandai ditangani.",
    ], numbered=True)
    add_callout(doc, "Insiden bukan penalti", "Catatan insiden tidak otomatis mengurangi nilai. Admin meninjau fakta dan membuat penalti terpisah bila memang diputuskan.", "info")

    add_heading(doc, "6. Koneksi dan pembaruan data", 1)
    add_list(doc, [
        "Halaman operasional menggunakan Realtime dan penyegaran berkala.",
        "Pilih Segarkan bila status dari petugas lain belum terlihat.",
        "Saat luring, jangan melakukan perubahan berulang karena mutasi memerlukan koneksi.",
        "Setelah jaringan pulih, muat ulang lalu periksa perubahan terakhir.",
        "Catat manual waktu kejadian jika aplikasi tidak dapat digunakan, lalu input setelah koneksi pulih sesuai arahan koordinator.",
    ])

    add_heading(doc, "7. Checklist Operator", 1)
    add_list(doc, [
        "Akun dan perangkat yang digunakan milik petugas yang bertugas.",
        "Nomor peserta dikonfirmasi sebelum setiap perubahan.",
        "Status dan waktu aktual sesuai kejadian lapangan.",
        "Titik atraksi pada layar sesuai penugasan.",
        "Setiap peserta pada titik atraksi memiliki status final yang jelas.",
        "Insiden penting sudah dicatat dan ditindaklanjuti.",
        "Tidak ada perubahan yang dilakukan memakai akun petugas lain.",
    ], checklist=True)
    add_shared_status_help(doc)
    return doc


def build_viewer():
    role = "Viewer Hasil"
    doc = new_document(role)
    add_cover(doc, role, "Membaca hasil resmi yang telah diterbitkan tanpa mengakses data internal.")

    add_heading(doc, "1. Membuka hasil", 1)
    add_body(doc, "Halaman hasil resmi tersedia pada rute /hasil. Halaman ini dapat dibuka tanpa login setelah panitia menerbitkan snapshot hasil. Akun dengan role Viewer juga diarahkan langsung ke halaman tersebut setelah login.")
    add_list(doc, [
        "Buka alamat aplikasi yang diberikan panitia, lalu pilih Lihat hasil karnaval yang diterbitkan, atau buka /hasil.",
        "Pilih tab Pendidikan atau Umum untuk melihat kategori yang diinginkan.",
        "Periksa label Published dan waktu penerbitan pada bagian bawah hasil.",
        "Gunakan tombol muat ulang browser bila panitia baru saja menerbitkan pembaruan.",
    ], numbered=True)
    add_callout(doc, "Belum diterbitkan", "Jika halaman menampilkan Hasil belum diterbitkan, data resmi belum tersedia. Viewer tidak dapat melihat preview atau hasil draf.", "info")

    add_heading(doc, "2. Membaca tabel hasil", 1)
    add_table(doc, ["Kolom", "Penjelasan"], [
        ("Peringkat", "Urutan peserta dalam kategori setelah seluruh aturan tie-break diterapkan."),
        ("No.", "Nomor urut peserta pada event."),
        ("Peserta", "Nama tim/lembaga dan tema penampilan."),
        ("Nilai Juri", "Nilai agregat dari lembar Juri yang sah; bukan rincian tiap Juri."),
        ("Atraksi", "Poin tambahan dari titik atraksi berstatus performed."),
        ("Penalti", "Jumlah pengurangan dari penalti confirmed."),
        ("Nilai akhir", "Nilai Juri agregat ditambah atraksi, dikurangi penalti, minimum 0."),
    ], [2200, 7160])
    add_body(doc, "Tiga kartu pemenang menampilkan Juara I, II, dan III untuk kategori yang sedang dipilih. Peringkat kategori Pendidikan dan Umum dihitung terpisah.")

    add_heading(doc, "3. Informasi yang tidak ditampilkan", 1)
    add_list(doc, [
        "Rincian nilai dari masing-masing Juri.",
        "Catatan Juri, catatan panitia, atau alasan internal.",
        "Hasil preview, hasil draf, dan lembar yang belum lengkap.",
        "Data akun, audit log, penugasan petugas, atau konfigurasi internal.",
    ])
    add_callout(doc, "Privasi dan integritas", "Jangan meminta kredensial petugas untuk melihat data internal. Halaman publik sengaja hanya menampilkan snapshot yang telah disetujui dan diterbitkan.", "warning")

    add_heading(doc, "4. Membagikan hasil", 1)
    add_list(doc, [
        "Bagikan tautan halaman /hasil, bukan tangkapan layar halaman internal.",
        "Sebutkan kategori dan waktu penerbitan bila membagikan hasil ke grup atau media sosial.",
        "Jika ada koreksi resmi, gunakan hasil yang terakhir diterbitkan oleh panitia.",
        "Untuk keberatan atau klarifikasi, hubungi kanal resmi panitia; jangan mengubah atau menyusun ulang tabel sebagai hasil resmi baru.",
    ])

    add_heading(doc, "5. Pemecahan masalah", 1)
    add_table(doc, ["Masalah", "Yang dilakukan"], [
        ("Hasil belum diterbitkan", "Tunggu pengumuman resmi panitia dan coba lagi kemudian."),
        ("Halaman lama masih tampil", "Lakukan hard refresh dengan Ctrl+Shift+R atau tutup dan buka kembali tab."),
        ("Kategori salah", "Pilih tab Pendidikan atau Umum pada bagian atas tabel."),
        ("Halaman tidak ditemukan", "Pastikan alamat berakhir dengan /hasil dan tidak ada salah ketik."),
        ("Data berbeda dari pengumuman", "Catat waktu penerbitan dan laporkan ke panitia tanpa menyebarkan dugaan."),
    ], [2900, 6460])
    add_heading(doc, "Checklist Viewer", 2)
    add_list(doc, [
        "Halaman menunjukkan status Published.",
        "Kategori yang dipilih sudah benar.",
        "Waktu penerbitan terlihat.",
        "Tautan yang dibagikan adalah /hasil.",
        "Tidak ada data internal atau rincian nilai tiap Juri pada halaman.",
    ], checklist=True)
    return doc


def audit_document(path):
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    attr = lambda name: f"{{{ns['w']}}}{name}"
    with ZipFile(path) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))
        styles_xml = ET.fromstring(archive.read("word/styles.xml"))
        numbering_xml = ET.fromstring(archive.read("word/numbering.xml"))

    section = document_xml.find(".//w:sectPr", ns)
    page_size = section.find("w:pgSz", ns)
    page_margins = section.find("w:pgMar", ns)
    assert page_size.get(attr("w")) == str(PAGE_WIDTH_DXA)
    assert page_size.get(attr("h")) == str(PAGE_HEIGHT_DXA)
    for side in ("top", "right", "bottom", "left"):
        assert page_margins.get(attr(side)) == "1440"
    assert page_margins.get(attr("header")) in {"708", "709"}
    assert page_margins.get(attr("footer")) in {"708", "709"}

    normal = styles_xml.find(".//w:style[@w:styleId='Normal']", ns)
    assert normal is not None
    assert normal.find("w:rPr/w:sz", ns).get(attr("val")) == "22"
    assert len(numbering_xml.findall(".//w:abstractNum", ns)) >= 3

    forbidden_prefixes = ("- ", "* ", "• ", "☐ ")
    for paragraph in document_xml.findall(".//w:body/w:p", ns):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns))
        assert not text.startswith(forbidden_prefixes), f"Fake list marker: {text}"
    assert audit_docx_tables(path) == 0


def save_document(document, filename):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    document.save(path)
    audit_document(path)
    print(path)
    return path


def main():
    outputs = [
        save_document(build_super_admin(), "Panduan-Super-Admin.docx"),
        save_document(build_admin(), "Panduan-Admin-Panitia.docx"),
        save_document(build_judge(), "Panduan-Juri.docx"),
        save_document(build_operator(), "Panduan-Operator-Lapangan.docx"),
        save_document(build_viewer(), "Panduan-Viewer-Hasil.docx"),
    ]
    print(f"Generated {len(outputs)} role guides in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
